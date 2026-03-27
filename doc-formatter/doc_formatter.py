#!/usr/bin/env python3
"""
doc-formatter - Fill Microsoft Word templates (.dotx) with text input using AI.
"""

import argparse
import io
import json
import os
import sys
import urllib.request
import urllib.error
import zipfile
from datetime import datetime
from pathlib import Path

try:
    import docx
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Environment ──────────────────────────────────────────────────────────────

def load_env():
    """Load .env from parent directory into environment variables."""
    env_path = Path(__file__).parent.parent / ".env"
    if not env_path.exists():
        return
    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            os.environ.setdefault(key.strip(), value.strip())


# ── Template handling ─────────────────────────────────────────────────────────

def open_dotx(path: Path) -> docx.Document:
    """Open a .dotx template as a python-docx Document.

    .dotx files differ from .docx only in the content type declaration.
    We patch that in-memory so python-docx can open them.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        b"wordprocessingml.template.main+xml",
                        b"wordprocessingml.document.main+xml",
                    )
                zout.writestr(item, data)
    buf.seek(0)
    return docx.Document(buf)


def extract_template_structure(doc: docx.Document) -> str:
    """Return a human-readable outline of the template structure."""
    lines = []

    # Collect all styles present in the template
    available_styles = sorted({p.style.name for p in doc.paragraphs} |
                               {p.style.name
                                for t in doc.tables
                                for row in t.rows
                                for cell in row.cells
                                for p in cell.paragraphs})
    lines.append(f"Available styles: {', '.join(available_styles)}")
    lines.append("")

    # Walk body elements in order (paragraphs and tables)
    body = doc.element.body
    table_index = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Find matching paragraph
            for p in doc.paragraphs:
                if p._element is child:
                    text = p.text.strip() or "(empty)"
                    lines.append(f"PARAGRAPH [{p.style.name}]: {text!r}")
                    break

        elif tag == "tbl":
            t = doc.tables[table_index]
            lines.append(f"TABLE {table_index}: {len(t.rows)} rows × {len(t.columns)} cols")
            for ri, row in enumerate(t.rows):
                seen_cells = set()
                row_parts = []
                for cell in row.cells:
                    cell_id = id(cell._element)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    text = " / ".join(p.text for p in cell.paragraphs if p.text) or "(empty)"
                    row_parts.append(text[:40])
                lines.append(f"  Row {ri}: {' | '.join(row_parts)}")
            table_index += 1

        elif tag == "sectPr":
            pass  # Section properties, skip

    return "\n".join(lines)


# ── Input handling ────────────────────────────────────────────────────────────

def read_input(input_arg: str) -> str:
    """Read input text from a file path or treat the argument as literal text."""
    path = Path(input_arg)
    if path.exists() and path.is_file():
        return path.read_text(encoding="utf-8")
    return input_arg


def generate_output_path(template_path: Path) -> Path:
    """Generate a timestamped output .docx path next to the template."""
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return template_path.parent / f"{template_path.stem}_{stamp}.docx"


# ── AI calls ──────────────────────────────────────────────────────────────────

SCHEMA_PROMPT = """Du bist ein Assistent, der Word-Dokument-Vorlagen analysiert.

Hier ist die Struktur einer Word-Vorlage:

{template_structure}

Analysiere diese Struktur und erstelle ein JSON-Schema, das beschreibt, welche Felder und Abschnitte befüllt werden müssen.

Gib **nur reines JSON** zurück (kein Markdown, kein Code-Block):
{{
  "title_style": "<style name of the title paragraph>",
  "fields": [
    {{
      "id": "<field identifier>",
      "label": "<label text as shown in template>",
      "location": "table:<table_index>:row:<row_index>:col:<col_index>",
      "type": "text|multiline"
    }}
  ],
  "sections": [
    {{
      "heading": "<section heading text>",
      "heading_style": "<style name>",
      "item_styles": ["<style1>", "<style2>"]
    }}
  ]
}}

Für "fields": erfasse alle Tabellenzellen, die Werte aufnehmen (die leeren Zellen neben Labels).
Für "sections": erfasse alle Abschnitte mit ihrem Stil und möglichen Unterstilen für Inhalte.
"""

FILL_PROMPT = """Du bist ein Assistent, der Word-Dokumente befüllt.

Hier ist das Schema der Vorlage:
{schema}

Hier ist der Eingabetext:
{input_text}

Befülle das Schema mit den passenden Inhalten aus dem Eingabetext.
Behalte die Struktur bei und nutze die vorhandenen Abschnitte.
Füge neue Abschnitte hinzu, falls nötig, mit dem passenden Stil.

Gib **nur reines JSON** zurück (kein Markdown, kein Code-Block):
{{
  "title": "<Dokumenttitel>",
  "fields": {{
    "<field_id>": "<value>"
  }},
  "sections": [
    {{
      "heading": "<Abschnittstitel>",
      "heading_style": "<style>",
      "items": [
        {{"style": "<style>", "text": "<Inhalt>"}}
      ]
    }}
  ]
}}

Hinweise:
- Für Teilnehmer-ähnliche Felder (type: multiline): trenne Einträge mit \\n
- Nutze die im Schema aufgelisteten Stile
- Befülle alle relevanten Felder aus dem Text
"""


def call_openrouter(api_key: str, prompt: str, model: str) -> str:
    """Call OpenRouter API and return the raw response content string."""
    payload = json.dumps({
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
    }).encode()

    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode(errors="replace")
        print(f"Error: OpenRouter HTTP {e.code}: {body[:300]}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: OpenRouter request failed: {e}", file=sys.stderr)
        sys.exit(1)

    return data["choices"][0]["message"]["content"].strip()


def parse_json_response(raw: str, context: str) -> dict:
    """Strip markdown fences from AI response and parse JSON."""
    if raw.startswith("```"):
        lines = raw.splitlines()
        raw = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"Error: Could not parse AI {context} response as JSON: {e}", file=sys.stderr)
        print(f"Raw response (first 400 chars):\n{raw[:400]}", file=sys.stderr)
        sys.exit(1)


def call_ai_schema(api_key: str, template_structure: str, model: str) -> dict:
    """Step 1: Ask AI to generate a fill schema from the template structure."""
    prompt = SCHEMA_PROMPT.format(template_structure=template_structure)
    raw = call_openrouter(api_key, prompt, model)
    return parse_json_response(raw, "schema")


def call_ai_fill(api_key: str, schema: dict, input_text: str, model: str) -> dict:
    """Step 2: Ask AI to fill the schema fields with content from input_text."""
    prompt = FILL_PROMPT.format(
        schema=json.dumps(schema, ensure_ascii=False, indent=2),
        input_text=input_text,
    )
    raw = call_openrouter(api_key, prompt, model)
    return parse_json_response(raw, "fill")


# ── Document filling ──────────────────────────────────────────────────────────

def _set_cell_text(cell, text: str):
    """Set the text of a table cell, handling multiline values."""
    lines = text.split("\n")
    # Clear existing runs in the first paragraph
    for run in cell.paragraphs[0].runs:
        run.text = ""
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = lines[0]
    else:
        cell.paragraphs[0].add_run(lines[0])
    # Add additional paragraphs for subsequent lines
    for line in lines[1:]:
        if line.strip():
            cell.add_paragraph(line)


def fill_document(template_path: Path, schema: dict, fill_data: dict, output_path: Path):
    """Apply AI-generated fill data to the template and save as output_path."""
    doc = open_dotx(template_path)
    body = doc.element.body

    # ── A: Fill title ────────────────────────────────────────────────────────
    title_style = schema.get("title_style", "Title")
    new_title = fill_data.get("title", "")
    if new_title:
        for p in doc.paragraphs:
            if p.style.name == title_style and p.text.strip():
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_title
                else:
                    p.add_run(new_title)
                break

    # ── B: Fill info table fields ────────────────────────────────────────────
    filled_fields = fill_data.get("fields", {})
    schema_fields = {f["id"]: f for f in schema.get("fields", [])}

    for field_id, value in filled_fields.items():
        if not value:
            continue
        field_def = schema_fields.get(field_id)
        if not field_def:
            continue
        loc = field_def.get("location", "")
        # Parse "table:0:row:1:col:3"
        parts = loc.split(":")
        if len(parts) == 6 and parts[0] == "table":
            try:
                ti, ri, ci = int(parts[1]), int(parts[3]), int(parts[5])
                if ti < len(doc.tables):
                    table = doc.tables[ti]
                    if ri < len(table.rows):
                        row = table.rows[ri]
                        # Deduplicate cells (merged cells appear multiple times)
                        seen = set()
                        unique_cells = []
                        for cell in row.cells:
                            cid = id(cell._element)
                            if cid not in seen:
                                seen.add(cid)
                                unique_cells.append(cell)
                        if ci < len(unique_cells):
                            _set_cell_text(unique_cells[ci], str(value))
            except (ValueError, IndexError):
                print(f"Warning: Could not locate field '{field_id}' at {loc}", file=sys.stderr)

    # ── C: Replace content sections ──────────────────────────────────────────
    # Find the first section heading paragraph element
    section_heading_style = "Überschrift 1.1"
    first_section_elem = None
    for p in doc.paragraphs:
        if p.style.name == section_heading_style:
            first_section_elem = p._element
            break

    if first_section_elem is not None:
        # Collect all body children from first section heading onward (except sectPr)
        elems_to_remove = []
        found = False
        for child in list(body):
            if child is first_section_elem:
                found = True
            if found:
                if child.tag.split("}")[-1] == "sectPr":
                    break
                elems_to_remove.append(child)
        for elem in elems_to_remove:
            body.remove(elem)

    # Add AI-generated sections
    sections = fill_data.get("sections", [])
    for section in sections:
        heading_style = section.get("heading_style", section_heading_style)
        # Validate style exists
        try:
            doc.styles[heading_style]
        except KeyError:
            heading_style = section_heading_style

        doc.add_paragraph(section.get("heading", ""), style=heading_style)

        for item in section.get("items", []):
            item_style = item.get("style", "Body Text")
            try:
                doc.styles[item_style]
            except KeyError:
                print(f"Warning: Style '{item_style}' not found, using 'Body Text'", file=sys.stderr)
                item_style = "Body Text"
            doc.add_paragraph(item.get("text", ""), style=item_style)

    doc.save(str(output_path))


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    load_env()

    parser = argparse.ArgumentParser(
        prog="doc-formatter",
        description="Fill a Word template (.dotx) with text input using AI.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Examples:
  doc-formatter Protokoll.dotx notes.txt
  doc-formatter Protokoll.dotx "Meeting on 2026-03-19..." --output result.docx
  doc-formatter Bericht.dotx bericht.md --model anthropic/claude-3-5-haiku
""",
    )
    parser.add_argument("template", help="Path to .dotx Word template file")
    parser.add_argument(
        "input",
        help="Path to input text file, or literal text string",
    )
    parser.add_argument(
        "--output", "-o",
        help="Output .docx file path (default: auto-generated with timestamp)",
    )
    parser.add_argument(
        "--model",
        default=os.environ.get("OPENROUTER_MODEL", "moonshotai/kimi-k2.5"),
        help="OpenRouter model ID (default: moonshotai/kimi-k2.5)",
    )

    args = parser.parse_args()

    # Validate template
    template_path = Path(args.template).resolve()
    if not template_path.exists():
        print(f"Error: Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    # Read input
    input_text = read_input(args.input)
    if not input_text.strip():
        print("Error: Input text is empty.", file=sys.stderr)
        sys.exit(1)

    # Determine output path
    output_path = Path(args.output).resolve() if args.output else generate_output_path(template_path)

    # Check API key
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        print("Error: OPENROUTER_API_KEY not set. Add it to the .env file or environment.", file=sys.stderr)
        sys.exit(1)

    print(f"Template:  {template_path}", file=sys.stderr)
    print(f"Input:     {len(input_text)} characters", file=sys.stderr)
    print(f"Output:    {output_path}", file=sys.stderr)
    print(f"Model:     {args.model}", file=sys.stderr)

    # Open template and extract structure
    print("\nExtracting template structure...", file=sys.stderr)
    try:
        doc = open_dotx(template_path)
    except Exception as e:
        print(f"Error: Could not open template: {e}", file=sys.stderr)
        sys.exit(1)
    template_structure = extract_template_structure(doc)

    # Step 1: Generate fill schema
    print("Step 1/2: Generating fill schema from template...", file=sys.stderr)
    schema = call_ai_schema(api_key, template_structure, args.model)
    print(f"  Schema: {len(schema.get('fields', []))} fields, {len(schema.get('sections', []))} sections", file=sys.stderr)

    # Step 2: Fill schema with input content
    print("Step 2/2: Filling template with input content...", file=sys.stderr)
    fill_data = call_ai_fill(api_key, schema, input_text, args.model)
    print(f"  Content: title={fill_data.get('title', '')!r}, {len(fill_data.get('sections', []))} sections", file=sys.stderr)

    # Fill and save document
    print("\nWriting output document...", file=sys.stderr)
    try:
        fill_document(template_path, schema, fill_data, output_path)
    except Exception as e:
        print(f"Error: Could not write document: {e}", file=sys.stderr)
        sys.exit(1)

    print(f"\nDone. Output saved to: {output_path}")
    sys.exit(0)


if __name__ == "__main__":
    main()
